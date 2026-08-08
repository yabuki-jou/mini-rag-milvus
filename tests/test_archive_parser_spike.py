"""验证智慧档案 V1 Parser 的定位与快照候选契约。"""

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from app.core.errors import AppError
from app.services.archive_parser_service import (
    MAX_PARSED_TEXT_CHARACTERS,
    MIN_EFFECTIVE_TEXT_CHARACTERS,
    ArchiveLocationType,
    parse_archive_document,
)


def _write_text_pdf(file_path: Path, text: str) -> None:
    """生成只含 ASCII 文本的最小 PDF，供 pypdf 真实提取验证使用。"""
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    stream = DecodedStreamObject()
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(
        f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode("latin-1")
    )
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(stream)
    with file_path.open("wb") as pdf_file:
        writer.write(pdf_file)


def test_text_and_markdown_normalize_newlines_and_keep_source_line_numbers(
    tmp_path: Path,
) -> None:
    """TXT/MD 必须固定 CRLF/CR 换行并保留未压缩的原始行号。"""
    text_file = tmp_path / "construction-plan.md"
    text_file.write_bytes(
        "# 施工方案\r\n\r\n编制单位：示例建设公司\r版本：V1.0\n".encode("utf-8")
    )

    result = parse_archive_document(text_file)

    assert [fragment.location_start for fragment in result.fragments] == [1, 3, 4]
    assert all(
        fragment.location_type == ArchiveLocationType.TEXT_LINE_RANGE
        for fragment in result.fragments
    )
    assert [fragment.content for fragment in result.fragments] == [
        "# 施工方案",
        "编制单位：示例建设公司",
        "版本：V1.0",
    ]


def test_docx_uses_nonempty_logical_blocks_and_table_rows_as_paragraphs(
    tmp_path: Path,
) -> None:
    """DOCX 的空段落不编号，列表文本和表格行按正文出现顺序编号。"""
    docx_file = tmp_path / "design-description.docx"
    document = DocxDocument()
    document.add_paragraph("设计说明")
    document.add_paragraph("")
    document.add_paragraph("1. 结构设计采用钢筋混凝土框架。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "编制单位"
    table.cell(0, 1).text = "示例设计院"
    table.cell(1, 0).text = "版本号"
    table.cell(1, 1).text = "V1.0"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "版本号：V1.0"
    document.save(docx_file)

    result = parse_archive_document(docx_file)

    assert [fragment.location_start for fragment in result.fragments] == [1, 2, 3, 4]
    assert [fragment.content for fragment in result.fragments] == [
        "设计说明",
        "1. 结构设计采用钢筋混凝土框架。",
        "编制单位 | 示例设计院",
        "版本号：V1.0",
    ]
    assert result.fragments[2].anchor_text == "编制单位 | 示例设计院"
    assert all(
        fragment.location_type == ArchiveLocationType.DOCX_PARAGRAPH
        for fragment in result.fragments
    )


def test_text_pdf_keeps_one_based_page_location(tmp_path: Path) -> None:
    """文本型 PDF 的真实提取结果必须保持原始页码。"""
    pdf_file = tmp_path / "project-contract.pdf"
    _write_text_pdf(pdf_file, "Project contract signed by Example Construction Company.")

    result = parse_archive_document(pdf_file)

    assert len(result.fragments) == 1
    assert result.fragments[0].location_type == ArchiveLocationType.PDF_PAGE
    assert result.fragments[0].location_start == 1
    assert "Project contract" in result.fragments[0].content


def test_blank_pdf_is_classified_as_scanned_pdf(tmp_path: Path) -> None:
    """没有可提取文本的 PDF 不能伪装为解析成功。"""
    pdf_file = tmp_path / "scanned.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with pdf_file.open("wb") as pdf_handle:
        writer.write(pdf_handle)

    with pytest.raises(AppError) as exc_info:
        parse_archive_document(pdf_file)

    assert exc_info.value.code == "SCANNED_PDF_UNSUPPORTED"
    assert exc_info.value.message == "暂不支持扫描件，请上传文本版资料。"


def test_short_text_pdf_is_not_misclassified_as_a_scanned_pdf(
    tmp_path: Path,
) -> None:
    """可提取文字但低于阈值的 PDF 应报告有效文本不足，而非扫描件。"""
    pdf_file = tmp_path / "short-text.pdf"
    _write_text_pdf(pdf_file, "Short text")

    with pytest.raises(AppError) as exc_info:
        parse_archive_document(pdf_file)

    assert exc_info.value.code == "PARSE_TEXT_UNAVAILABLE"


def test_repeated_parse_of_same_file_has_same_fragment_order_and_snapshot_hash(
    tmp_path: Path,
) -> None:
    """同一原文件和固定归一化规则必须生成完全相同的快照。"""
    text_file = tmp_path / "meeting.txt"
    text_file.write_text(
        "项目会议纪要\n会议决定按计划提交验收资料。\n",
        encoding="utf-8",
    )

    first_result = parse_archive_document(text_file)
    second_result = parse_archive_document(text_file)

    assert first_result.fragments == second_result.fragments
    assert first_result.snapshot_hash == second_result.snapshot_hash
    assert len(first_result.snapshot_hash) == 64


def test_effective_text_threshold_rejects_shell_text(tmp_path: Path) -> None:
    """少于 20 个归一化字符的文本不满足归档所需的有效文本条件。"""
    text_file = tmp_path / "shell.txt"
    text_file.write_text("文档标题", encoding="utf-8")

    with pytest.raises(AppError) as exc_info:
        parse_archive_document(text_file)

    assert MIN_EFFECTIVE_TEXT_CHARACTERS == 20
    assert exc_info.value.code == "PARSE_TEXT_UNAVAILABLE"


def test_effective_text_limit_rejects_oversized_snapshot(tmp_path: Path) -> None:
    """快照必须有明确上限，避免将无限文本写入后续存储。"""
    text_file = tmp_path / "oversized.txt"
    text_file.write_text("a" * (MAX_PARSED_TEXT_CHARACTERS + 1), encoding="utf-8")

    with pytest.raises(AppError) as exc_info:
        parse_archive_document(text_file)

    assert MAX_PARSED_TEXT_CHARACTERS == 1_000_000
    assert exc_info.value.code == "PARSE_TEXT_TOO_LARGE"
