"""生成智慧档案 V1 的虚构验收资料与人工标注。"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import shutil
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen.canvas import Canvas


PROJECT_ROOT = Path(__file__).resolve().parents[1]
# 该目录仅用于本地查看与测试；由 .gitignore 排除，不作为仓库交付物。
DATASET_ROOT = PROJECT_ROOT / "tests" / "pytest_docs"
DOCUMENTS_ROOT = DATASET_ROOT / "documents"
LABELS_ROOT = DATASET_ROOT / "labels"
README_PATH = DATASET_ROOT / "README.md"
PDF_FONT_NAME = "STSong-Light"

FIELD_LABELS = {
    "TITLE": "资料标题",
    "DOCUMENT_TYPE": "资料类型",
    "DOCUMENT_DATE": "文档日期",
    "AUTHORING_ORGANIZATION": "编制单位",
    "VERSION_NUMBER": "版本号",
    "PROJECT_STAGE": "项目阶段",
    "KEYWORDS": "关键词",
}


def _record(
    record_id: str,
    project_id: str,
    filename: str,
    source_format: str,
    document_type: str,
    project_stage: str,
    title: str,
    date: str | None,
    organization: str | None,
    version: str | None,
    keywords: list[str] | None,
    body: list[str],
    evidence_points: dict[str, str],
) -> dict[str, Any]:
    """构造一份正常验收资料的固定元数据。"""
    return {
        "id": record_id,
        "project_id": project_id,
        "filename": filename,
        "source_format": source_format,
        "document_type": document_type,
        "project_stage": project_stage,
        "title": title,
        "date": date,
        "organization": organization,
        "version": version,
        "keywords": keywords,
        "body": body,
        "evidence_points": evidence_points,
    }


NORMAL_RECORDS = [
    _record(
        "A-01",
        "alpha",
        "alpha_contract.pdf",
        "PDF",
        "CONTRACT",
        "PREPARATION",
        "星河办公楼改造工程施工总承包合同",
        "2025-03-18",
        "星河建设有限公司",
        "V1.0",
        ["施工总承包", "履约保证", "项目合同"],
        [
            "合同双方：星河建设有限公司与星河资产管理有限公司。",
            "本合同不包含付款比例的约定。",
        ],
        {},
    ),
    _record(
        "A-02",
        "alpha",
        "alpha_design_description.docx",
        "DOCX",
        "DESIGN",
        "DESIGN",
        "星河办公楼改造工程设计说明",
        "2025-04-08",
        "北辰设计院",
        "V2.1",
        ["结构设计", "办公楼改造", "设计说明"],
        ["本说明适用于星河办公楼改造工程的结构与给排水设计。"],
        {},
    ),
    _record(
        "A-03",
        "alpha",
        "alpha_construction_plan.txt",
        "TXT",
        "CONSTRUCTION",
        "CONSTRUCTION",
        "星河办公楼改造工程施工方案",
        "2025-05-12",
        "星河建设有限公司",
        "V1.2",
        ["施工方案", "安全交底", "改造工程"],
        ["本方案要求进场人员完成安全交底后方可进行拆改施工。"],
        {},
    ),
    _record(
        "A-04",
        "alpha",
        "alpha_meeting_minutes.md",
        "MD",
        "MEETING_MINUTES",
        "CROSS_STAGE",
        "星河办公楼改造工程第一次协调会议纪要",
        "2025-05-20",
        "星河项目部",
        None,
        ["协调会议", "设计变更", "进度管理"],
        [
            "会议决定：北辰设计院在 3 个工作日内提交设计变更说明。",
            "会议要求施工单位在变更确认后更新施工组织安排。",
        ],
        {"design_change_decision": "会议决定：北辰设计院在 3 个工作日内提交设计变更说明。"},
    ),
    _record(
        "A-05",
        "alpha",
        "alpha_acceptance_report.pdf",
        "PDF",
        "ACCEPTANCE",
        "ACCEPTANCE",
        "星河办公楼改造工程竣工验收报告",
        "2025-09-30",
        "星河项目验收组",
        "V1.0",
        ["竣工验收", "资料齐全", "备案"],
        ["验收结论：资料齐全，建议进入项目竣工验收备案。"],
        {"acceptance_conclusion": "验收结论：资料齐全，建议进入项目竣工验收备案。"},
    ),
    _record(
        "A-06",
        "alpha",
        "alpha_site_photo_record.md",
        "MD",
        "OTHER",
        "CONSTRUCTION",
        "星河办公楼改造工程现场影像记录",
        "2025-06-03",
        "星河项目部",
        None,
        ["现场影像", "施工记录"],
        ["本记录用于索引施工阶段的现场影像，不替代施工方案或验收报告。"],
        {},
    ),
    _record(
        "B-01",
        "beta",
        "beta_equipment_contract.docx",
        "DOCX",
        "CONTRACT",
        "PREPARATION",
        "云港仓储中心设备采购合同",
        "2025-02-26",
        "云港仓储建设有限公司",
        "V1.0",
        ["设备采购", "仓储中心", "供货合同"],
        ["供货单位：云港设备供应有限公司。", "本合同约定货物到场后进行数量核验。"],
        {"supplier": "供货单位：云港设备供应有限公司。"},
    ),
    _record(
        "B-02",
        "beta",
        "beta_design_spec.md",
        "MD",
        "DESIGN",
        "DESIGN",
        "云港仓储中心设计技术说明",
        "2025-03-15",
        "海岳建筑设计事务所",
        "V3.0",
        ["仓储设计", "消防分区", "技术说明"],
        ["本说明明确仓储区按消防分区要求布置，并与设备清单对应。"],
        {},
    ),
    _record(
        "B-03",
        "beta",
        "beta_construction_plan.pdf",
        "PDF",
        "CONSTRUCTION",
        "CONSTRUCTION",
        "云港仓储中心消防施工方案",
        "2025-04-22",
        "云港安装工程有限公司",
        "V1.1",
        ["消防施工", "仓储中心", "安全交底"],
        ["施工要求：消防管线安装完成后应组织隐蔽工程检查。"],
        {},
    ),
    _record(
        "B-04",
        "beta",
        "beta_meeting_minutes.docx",
        "DOCX",
        "MEETING_MINUTES",
        "CROSS_STAGE",
        "云港仓储中心设备进场协调会纪要",
        "2025-05-08",
        "云港项目部",
        "V1.0",
        ["设备进场", "协调会议", "仓储中心"],
        ["会议决定：设备供应单位于 5 月 18 日前提交到货计划。"],
        {},
    ),
    _record(
        "B-05",
        "beta",
        "beta_acceptance_report.txt",
        "TXT",
        "ACCEPTANCE",
        "ACCEPTANCE",
        "云港仓储中心消防验收报告",
        "2025-10-12",
        "云港项目验收组",
        "V1.0",
        ["消防验收", "现场抽查", "仓储中心"],
        ["消防验收结论：消防验收资料齐全，现场抽查结果符合要求。"],
        {"fire_conclusion": "消防验收结论：消防验收资料齐全，现场抽查结果符合要求。"},
    ),
    _record(
        "B-06",
        "beta",
        "beta_equipment_delivery.txt",
        "TXT",
        "OTHER",
        "CONSTRUCTION",
        "云港仓储中心设备到货记录",
        "2025-05-21",
        "云港项目物资组",
        None,
        ["设备到货", "物资记录"],
        ["本记录用于登记设备到货数量，不替代采购合同或验收报告。"],
        {},
    ),
]


def _metadata_lines(record: dict[str, Any]) -> list[str]:
    """按固定顺序建立可被字段证据直接引用的原文行。"""
    lines = ["工程项目资料", f"资料标题：{record['title']}"]
    type_name = {
        "CONTRACT": "合同资料",
        "DESIGN": "设计资料",
        "CONSTRUCTION": "施工资料",
        "MEETING_MINUTES": "会议纪要",
        "ACCEPTANCE": "验收资料",
        "OTHER": "其他资料",
    }[record["document_type"]]
    stage_name = {
        "PREPARATION": "前期准备",
        "DESIGN": "设计阶段",
        "CONSTRUCTION": "施工阶段",
        "ACCEPTANCE": "验收阶段",
        "CROSS_STAGE": "跨阶段",
        "OTHER_STAGE": "其他阶段",
    }[record["project_stage"]]
    lines.append(f"资料类型：{type_name}")
    if record["date"]:
        lines.append(f"文档日期：{record['date']}")
    if record["organization"]:
        lines.append(f"编制单位：{record['organization']}")
    if record["version"]:
        lines.append(f"版本号：{record['version']}")
    lines.append(f"项目阶段：{stage_name}")
    if record["keywords"]:
        lines.append(f"关键词：{'；'.join(record['keywords'])}")
    return lines + record["body"]


def _set_docx_font(run: Any, size: int, bold: bool = False) -> None:
    """显式写入中文字体，避免 Word/LibreOffice 使用不一致的默认字体。"""
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def _write_docx(file_path: Path, lines: list[str]) -> None:
    """生成结构简单、可视觉检查的单页虚构 DOCX。"""
    document = DocxDocument()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for index, line in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            _set_docx_font(run, 16, bold=True)
            run.font.color.rgb = RGBColor(46, 116, 181)
        else:
            run = paragraph.add_run(line)
            _set_docx_font(run, 11)
    document.save(file_path)


def _write_pdf(file_path: Path, lines: list[str], scanned: bool = False) -> None:
    """生成带可提取中文文本的 PDF，或用于扫描件边界的无文本 PDF。"""
    pdfmetrics.registerFont(UnicodeCIDFont(PDF_FONT_NAME))
    canvas = Canvas(str(file_path), pagesize=letter)
    if scanned:
        canvas.setFillGray(0.85)
        canvas.rect(72, 420, 468, 210, fill=1, stroke=0)
        canvas.setFillGray(0.6)
        canvas.rect(96, 450, 420, 18, fill=1, stroke=0)
        canvas.rect(96, 485, 380, 18, fill=1, stroke=0)
        canvas.save()
        return

    text = canvas.beginText(72, 720)
    text.setFont(PDF_FONT_NAME, 12)
    for index, line in enumerate(lines):
        if index == 0:
            text.setFont(PDF_FONT_NAME, 16)
            text.textLine(line)
            text.setFont(PDF_FONT_NAME, 12)
        else:
            text.textLine(line)
        text.moveCursor(0, -5)
    canvas.drawText(text)
    canvas.save()


def _write_text(file_path: Path, lines: list[str], markdown: bool) -> None:
    """以固定 LF 和 UTF-8 写入 TXT/MD，保证 P01 行号规则可复现。"""
    rendered_lines = lines.copy()
    if markdown:
        rendered_lines[0] = f"# {rendered_lines[0]}"
    file_path.write_text("\n".join(rendered_lines) + "\n", encoding="utf-8")


def _file_sha256(file_path: Path) -> str:
    """返回原文件哈希，供重复上传场景和资料可追溯性验证使用。"""
    return hashlib.sha256(file_path.read_bytes()).hexdigest()


def _location_for_line(source_format: str, line_number: int) -> tuple[str, int, int]:
    """将资料行位置转换为已冻结的证据定位类型。"""
    if source_format == "PDF":
        return "PDF_PAGE", 1, 1
    if source_format == "DOCX":
        return "DOCX_PARAGRAPH", line_number, line_number
    return "TEXT_LINE_RANGE", line_number, line_number


def _build_label(record: dict[str, Any], file_hash: str) -> dict[str, Any]:
    """从同一份原文行构建人工字段 Ground Truth，避免手工标注漂移。"""
    lines = _metadata_lines(record)
    fields: dict[str, Any] = {}
    source_values = {
        "TITLE": record["title"],
        "DOCUMENT_TYPE": record["document_type"],
        "DOCUMENT_DATE": record["date"],
        "AUTHORING_ORGANIZATION": record["organization"],
        "VERSION_NUMBER": record["version"],
        "PROJECT_STAGE": record["project_stage"],
        "KEYWORDS": record["keywords"],
    }
    for field_name, value in source_values.items():
        if value is None:
            continue
        label = FIELD_LABELS[field_name]
        evidence_line = next(
            line_number
            for line_number, line in enumerate(lines, start=1)
            if line.startswith(f"{label}：")
        )
        location_type, location_start, location_end = _location_for_line(
            record["source_format"], evidence_line
        )
        fields[field_name] = {
            "value": value,
            "evidence": [
                {
                    "location_type": location_type,
                    "location_start": location_start,
                    "location_end": location_end,
                    "excerpt": lines[evidence_line - 1],
                }
            ],
        }

    evidence_points: dict[str, Any] = {}
    for point_name, excerpt in record["evidence_points"].items():
        line_number = lines.index(excerpt) + 1
        location_type, location_start, location_end = _location_for_line(
            record["source_format"], line_number
        )
        evidence_points[point_name] = {
            "location_type": location_type,
            "location_start": location_start,
            "location_end": location_end,
            "excerpt": excerpt,
        }

    return {
        "id": record["id"],
        "project_id": record["project_id"],
        "relative_path": f"documents/{record['filename']}",
        "source_format": record["source_format"],
        "scenario": "normal",
        "file_sha256": file_hash,
        "expected_document_type": record["document_type"],
        "expected_project_stage": record["project_stage"],
        "expected_fields": fields,
        "evidence_points": evidence_points,
    }


def _question(
    question_id: str,
    category: str,
    project_id: str,
    question: str,
    expected_answer: str | None,
    source_document_id: str | None = None,
    source_field: str | None = None,
    source_point: str | None = None,
) -> dict[str, Any]:
    """构造固定问答用例，证据在生成后由文档标注引用。"""
    return {
        "id": question_id,
        "category": category,
        "project_id": project_id,
        "question": question,
        "expected_answer": expected_answer,
        "source_document_id": source_document_id,
        "source_field": source_field,
        "source_point": source_point,
    }


QUESTION_SEEDS = [
    _question("Q-01", "GROUNDED", "alpha", "星河项目施工总承包合同的签订日期是多少？", "2025-03-18", "A-01", "DOCUMENT_DATE"),
    _question("Q-02", "GROUNDED", "alpha", "星河项目设计说明的编制单位是什么？", "北辰设计院", "A-02", "AUTHORING_ORGANIZATION"),
    _question("Q-03", "GROUNDED", "alpha", "星河项目施工方案适用哪个项目阶段？", "施工阶段", "A-03", "PROJECT_STAGE"),
    _question("Q-04", "GROUNDED", "alpha", "第一次协调会议对设计变更作出了什么决定？", "北辰设计院在 3 个工作日内提交设计变更说明。", "A-04", source_point="design_change_decision"),
    _question("Q-05", "GROUNDED", "alpha", "星河项目竣工验收报告的结论是什么？", "资料齐全，建议进入项目竣工验收备案。", "A-05", source_point="acceptance_conclusion"),
    _question("Q-06", "GROUNDED", "beta", "云港仓储中心设备采购合同的供货单位是什么？", "云港设备供应有限公司", "B-01", source_point="supplier"),
    _question("Q-07", "GROUNDED", "beta", "云港仓储中心设计技术说明的版本号是什么？", "V3.0", "B-02", "VERSION_NUMBER"),
    _question("Q-08", "GROUNDED", "beta", "云港仓储中心消防验收报告的结论是什么？", "消防验收资料齐全，现场抽查结果符合要求。", "B-05", source_point="fire_conclusion"),
    _question("Q-09", "NO_EVIDENCE", "alpha", "星河施工总承包合同约定的付款比例是多少？", None),
    _question("Q-10", "NO_EVIDENCE", "beta", "云港仓储中心的实际开工日期是什么？", None),
    _question("Q-11", "ISOLATION", "alpha", "云港仓储中心消防验收报告的结论是什么？", None, "B-05", source_point="fire_conclusion"),
    _question("Q-12", "ISOLATION", "beta", "星河项目设计说明的编制单位是什么？", None, "A-02", "AUTHORING_ORGANIZATION"),
]


def _build_questions(document_labels: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    """把问题与已有的字段/正文 Ground Truth 证据连接起来。"""
    questions = []
    for seed in QUESTION_SEEDS:
        question = seed.copy()
        document_id = question.pop("source_document_id")
        field_name = question.pop("source_field")
        point_name = question.pop("source_point")
        if document_id and question["category"] == "GROUNDED":
            label = document_labels[document_id]
            evidence = (
                label["expected_fields"][field_name]["evidence"]
                if field_name
                else [label["evidence_points"][point_name]]
            )
            question["expected_evidence"] = {
                "document_id": document_id,
                "relative_path": label["relative_path"],
                "items": evidence,
            }
        elif question["category"] == "ISOLATION":
            question["hidden_evidence_in_other_project"] = document_id
        questions.append(question)
    return questions


def _write_record(record: dict[str, Any]) -> Path:
    """按资料格式生成一份正常虚构原文件。"""
    output_path = DOCUMENTS_ROOT / record["filename"]
    lines = _metadata_lines(record)
    if record["source_format"] == "PDF":
        _write_pdf(output_path, lines)
    elif record["source_format"] == "DOCX":
        _write_docx(output_path, lines)
    else:
        _write_text(output_path, lines, markdown=record["source_format"] == "MD")
    return output_path


def _write_anomalies(document_labels: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """创建字段缺失/分类混淆、重复和扫描件三个非计分异常样本。"""
    ambiguous_path = DOCUMENTS_ROOT / "anomaly_ambiguous_missing_fields.md"
    ambiguous_lines = [
        "# 现场协调记录",
        "记录内容：设计变更尚未确认，施工班组暂缓相关工序。",
        "备注：请项目组后续人工判断资料类型和项目阶段。",
    ]
    ambiguous_path.write_text("\n".join(ambiguous_lines) + "\n", encoding="utf-8")

    original = next(label for label in document_labels if label["id"] == "A-03")
    duplicate_path = DOCUMENTS_ROOT / "anomaly_duplicate_alpha_construction_plan.txt"
    shutil.copyfile(DOCUMENTS_ROOT / Path(original["relative_path"]).name, duplicate_path)

    scanned_path = DOCUMENTS_ROOT / "anomaly_scanned.pdf"
    _write_pdf(scanned_path, [], scanned=True)

    return [
        {
            "id": "X-01",
            "project_id": "alpha",
            "relative_path": f"documents/{ambiguous_path.name}",
            "source_format": "MD",
            "scenario": "MISSING_FIELDS_AND_AMBIGUOUS_CLASSIFICATION",
            "file_sha256": _file_sha256(ambiguous_path),
            "expected_document_type": "OTHER",
            "identifiable_fields": ["TITLE"],
            "note": "不含明确资料类型、日期、编制单位、版本号或阶段，人工应决定为 OTHER 或更正。",
        },
        {
            "id": "X-02",
            "project_id": "alpha",
            "relative_path": f"documents/{duplicate_path.name}",
            "source_format": "TXT",
            "scenario": "EXACT_FILE_HASH_DUPLICATE",
            "file_sha256": _file_sha256(duplicate_path),
            "duplicate_of": "A-03",
            "note": "与 A-03 字节完全相同，应触发当前项目内 file_hash 重复冲突。",
        },
        {
            "id": "X-03",
            "project_id": "beta",
            "relative_path": f"documents/{scanned_path.name}",
            "source_format": "PDF",
            "scenario": "SCANNED_PDF_NO_TEXT",
            "file_sha256": _file_sha256(scanned_path),
            "expected_error_code": "SCANNED_PDF_UNSUPPORTED",
            "note": "PDF 只含图形，不含可提取文本。",
        },
    ]


def _write_local_readme() -> None:
    """说明本地测试资料的用途，避免被误当成业务数据或 Git 交付物。"""
    lines = [
        "# 智慧档案 V1 本地虚构验收资料",
        "",
        "此目录仅供本地查看与 pytest 使用，已由 `.gitignore` 排除，不能上传 GitHub。",
        "",
        "- `documents/`：12 份正常虚构资料和 3 份异常资料。",
        "- `labels/document-ground-truth.json`：资料类型、字段和原文证据的人工标注。",
        "- `labels/question-ground-truth.json`：8 个有依据、2 个无依据、2 个项目隔离问答。",
        "",
        "所有内容均为虚构演示资料，不代表真实工程项目的法定或行业归档要求。",
        "不得替换为真实商业秘密、个人敏感信息或未获授权的资料。",
        "",
        "若目录不存在，`tests/test_archive_v1_evaluation_data.py` 会调用",
        "`scripts/generate_archive_v1_eval_data.py` 自动重建。",
    ]
    README_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    """生成资料、人工标注、问题集与原文件哈希清单。"""
    DOCUMENTS_ROOT.mkdir(parents=True, exist_ok=True)
    LABELS_ROOT.mkdir(parents=True, exist_ok=True)
    document_labels: list[dict[str, Any]] = []
    label_index: dict[str, dict[str, Any]] = {}

    for record in NORMAL_RECORDS:
        output_path = _write_record(record)
        label = _build_label(record, _file_sha256(output_path))
        document_labels.append(label)
        label_index[label["id"]] = label

    anomaly_labels = _write_anomalies(document_labels)
    manifest = {
        "dataset_id": "archive-v1-fictional-evaluation-v1",
        "scope": "虚构演示与验收资料，不代表真实工程项目的法定或行业归档要求。",
        "projects": [
            {"id": "alpha", "name": "星河办公楼改造工程"},
            {"id": "beta", "name": "云港仓储中心项目"},
        ],
        "normal_documents": document_labels,
        "anomaly_documents": anomaly_labels,
    }
    questions = {
        "dataset_id": manifest["dataset_id"],
        "questions": _build_questions(label_index),
    }
    (LABELS_ROOT / "document-ground-truth.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    (LABELS_ROOT / "question-ground-truth.json").write_text(
        json.dumps(questions, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    _write_local_readme()
    print(f"generated {len(document_labels)} normal and {len(anomaly_labels)} anomaly documents")


if __name__ == "__main__":
    main()
